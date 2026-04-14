import io
import asyncio
from datetime import datetime, date
from typing import List, Dict, Any, Tuple, Optional

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_JUSTIFY, TA_LEFT
from reportlab.lib.units import cm

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

MarkdownImportError = None
try:
    import markdown2
except ImportError as e:
    markdown2 = None
    import logging
    logging.warning("Markdown2 library not found. Markdown formatting will be basic. Please install with: pip install markdown2")

from .models import DocumentFormat
from .config_models.config_models import BenefitTypeInfo, BenefitTypeDocuments, DocumentDetail

def mask_personal_data(personal_data: Dict[str, Any]) -> Dict[str, Any]:
    """Маскирует чувствительные персональные данные."""
    masked_data = {
        "full_name": "[ФИО СКРЫТО]",
        "birth_date": "[ДАТА РОЖДЕНИЯ СКРЫТА]",
        "snils": "[СНИЛС СКРЫТ]",
        "gender": personal_data.get("gender", "[ПОЛ СКРЫТ]"),
        "citizenship": "[ГРАЖДАНСТВО СКРЫТО]",
        "name_change_info": {
            "old_full_name": "[ПРЕЖНЕЕ ФИО СКРЫТО]",
            "date_changed": "[ДАТА СМЕНЫ ФИО СКРЫТА]"
        } if personal_data and personal_data.get("name_change_info") else {},
        "dependents": "[КОЛИЧЕСТВО ИЖДИВЕНЦЕВ СКРЫТО]"
    }
    return masked_data

def _get_benefit_type_display_name(benefit_type_id: str, benefit_types_config: List[BenefitTypeInfo]) -> str:
    if benefit_types_config:
        for bt in benefit_types_config:
            if bt.id == benefit_type_id:
                return bt.display_name
    return benefit_type_id

def _get_document_display_name(doc_id: str, doc_requirements_config: Dict[str, BenefitTypeDocuments], benefit_type_id_for_context: Optional[str] = None) -> str:
    if benefit_type_id_for_context and doc_requirements_config and benefit_type_id_for_context in doc_requirements_config:
        for doc_detail in doc_requirements_config[benefit_type_id_for_context].documents:
            if doc_detail.id == doc_id:
                return doc_detail.name
    if doc_requirements_config:
        for bt_id, reqs in doc_requirements_config.items():
            for doc_detail in reqs.documents:
                if doc_detail.id == doc_id:
                    return doc_detail.name
    return doc_id

def _convert_markdown_to_html_for_reportlab(md_text: str) -> str:
    """Конвертирует Markdown в HTML, подходящий для ReportLab Paragraph."""
    import logging
    logger_services = logging.getLogger(__name__)
    logger_services.debug(f"Original Markdown for PDF: {md_text!r}")

    if markdown2:
        html = markdown2.markdown(md_text, extras=["break-on-newline", "cuddled-lists", "smarty-pants"])
        logger_services.debug(f"Converted HTML (using markdown2) for PDF: {html!r}")
        return html
    else:
        logger_services.warning("markdown2 not available, using basic HTML escaping for PDF.")
        import html as html_converter
        escaped_text = html_converter.escape(md_text)
        html_fallback = escaped_text.replace("\n", "<br/>")
        logger_services.debug(f"Fallback HTML for PDF: {html_fallback!r}")
        return html_fallback

def _strip_markdown_for_docx(md_text: str) -> str:
    """Удаляет Markdown разметку из текста (упрощенный вариант)."""
    if markdown2:
        html_text = markdown2.markdown(md_text)
        import re
        clean_text = re.sub(r'<[^>]+>', '', html_text)
        clean_text = clean_text.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
        return clean_text
    else:
        text = md_text.replace("**", "").replace("__", "")\
                      .replace("*", "").replace("_", "")\
                      .replace("\n", " ")
        return text

def _generate_pdf_report(
    case_details: Dict[str, Any],
    benefit_types_list_config: List[BenefitTypeInfo],
    doc_requirements_config: Dict[str, BenefitTypeDocuments]
) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=2*cm, rightMargin=1.5*cm,
                            topMargin=1.5*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    elements = []

    font_name = 'Helvetica'
    bold_font_name = 'Helvetica-Bold'
    
    styles['Normal'].fontName = font_name
    styles['Heading1'].fontName = bold_font_name
    styles['Heading2'].fontName = bold_font_name
    styles['Heading3'].fontName = bold_font_name
    styles['Bullet'].fontName = font_name
    styles['Definition'].fontName = font_name

    title_style = ParagraphStyle(name='TitleStyle', parent=styles['Heading1'], fontSize=16, alignment=TA_CENTER, spaceAfter=12, fontName=bold_font_name)
    doc_date_style = ParagraphStyle(name='DocDateStyle', parent=styles['Normal'], fontSize=10, alignment=TA_RIGHT, spaceAfter=10, fontName=font_name)
    system_name_style = ParagraphStyle(name='SystemNameStyle', parent=styles['Normal'], fontSize=10, alignment=TA_CENTER, spaceBefore=6, spaceAfter=12, fontName=font_name)
    
    section_title_style = ParagraphStyle(name='SectionTitle', parent=styles['Heading2'], fontSize=14, spaceBefore=12, spaceAfter=6, fontName=bold_font_name)
    subsection_title_style = ParagraphStyle(name='SubSectionTitle', parent=styles['Heading3'], fontSize=12, spaceBefore=8, spaceAfter=4, fontName=bold_font_name)
    
    normal_style_justify = ParagraphStyle(name='NormalJustify', parent=styles['Normal'], fontSize=12, alignment=TA_JUSTIFY, spaceAfter=6, fontName=font_name)
    normal_style_left = ParagraphStyle(name='NormalLeft', parent=styles['Normal'], fontSize=12, alignment=TA_LEFT, spaceAfter=6, fontName=font_name)
    label_style = ParagraphStyle(name='LabelStyle', parent=normal_style_left, fontSize=12, fontName=font_name)
    
    error_style = ParagraphStyle(name='ErrorStyle', parent=normal_style_justify, textColor=colors.red, spaceAfter=6, fontName=font_name)
    footer_style = ParagraphStyle(name='FooterStyle', parent=styles['Normal'], fontSize=10, alignment=TA_CENTER, spaceBefore=24, fontName=font_name)

    case_id = case_details.get("id", "N/A")
    current_date_str = datetime.now().strftime("%d.%m.%Y")
    
    personal_data = case_details.get("personal_data", {})
    masked_personal_data = mask_personal_data(personal_data if personal_data else {})

    elements.append(Paragraph(f"ЗАКЛЮЧЕНИЕ № {case_id}", title_style))
    elements.append(Paragraph(f"по запросу о мерах социальной поддержки", styles['Normal']))
    elements.append(Paragraph(f"Дата формирования: {current_date_str}", doc_date_style))
    elements.append(Spacer(1, 0.5*cm))

    elements.append(Paragraph("1. Сведения о заявителе (обезличенные)", section_title_style))
    elements.append(Paragraph(f"<b>Фамилия, имя, отчество (при наличии):</b> {masked_personal_data['full_name']}", label_style))
    elements.append(Paragraph(f"<b>Дата рождения:</b> {masked_personal_data['birth_date']}", label_style))
    elements.append(Paragraph(f"<b>СНИЛС:</b> {masked_personal_data['snils']}", label_style))
    elements.append(Paragraph(f"<b>Пол:</b> {masked_personal_data['gender']}", label_style))
    elements.append(Paragraph(f"<b>Гражданство:</b> {masked_personal_data['citizenship']}", label_style))
    if masked_personal_data.get("name_change_info") and masked_personal_data["name_change_info"].get('old_full_name') != "[ПРЕЖНЕЕ ФИО СКРЫТО]":
        elements.append(Paragraph("<b>Сведения о ранее измененном ФИО:</b>", label_style))
        elements.append(Paragraph(f"  Прежнее ФИО: {masked_personal_data['name_change_info']['old_full_name']}", label_style))
        elements.append(Paragraph(f"  Дата изменения: {masked_personal_data['name_change_info']['date_changed']}", label_style))
    elements.append(Paragraph(f"<b>Количество заявленных иждивенцев:</b> {masked_personal_data['dependents']}", label_style))
    elements.append(Spacer(1, 0.5*cm))

    elements.append(Paragraph("2. Запрашиваемый вид поддержки", section_title_style))
    benefit_type_id = case_details.get("pension_type", "Не указан")
    benefit_type_name = _get_benefit_type_display_name(benefit_type_id, benefit_types_list_config)
    elements.append(Paragraph(benefit_type_name, normal_style_justify))
    elements.append(Spacer(1, 0.5*cm))

    elements.append(Paragraph("3. Представленные сведения и документы", section_title_style))
    
    disability_info = case_details.get("disability")
    if disability_info:
        elements.append(Paragraph("3.1. Сведения об инвалидности", subsection_title_style))
        elements.append(Paragraph(f"<b>Группа инвалидности:</b> {disability_info.get('group', 'Не указана')}", label_style))
        dis_date = disability_info.get('date')
        dis_date_str = dis_date.strftime("%d.%m.%Y") if isinstance(dis_date, datetime) or isinstance(dis_date, date) else str(dis_date or 'Не указана')
        elements.append(Paragraph(f"<b>Дата установления:</b> {dis_date_str}", label_style))
        elements.append(Paragraph(f"<b>Номер справки МСЭ:</b> {disability_info.get('cert_number', 'Не указан')}", label_style))

    work_experience_info = case_details.get("work_experience")
    pension_points = case_details.get("pension_points")
    if work_experience_info or pension_points is not None:
        elements.append(Paragraph("3.2. Сведения о трудовом стаже", subsection_title_style))
        if work_experience_info:
            total_years = work_experience_info.get('calculated_total_years', 'Не указан')
            elements.append(Paragraph(f"<b>Общий страховой стаж (лет):</b> {total_years}", label_style))
        if pension_points is not None:
            elements.append(Paragraph(f"<b>Индивидуальный пенсионный коэффициент (ИПК):</b> {pension_points}", label_style))
        
        if work_experience_info and work_experience_info.get("records"):
            elements.append(Paragraph("<b>Периоды трудовой деятельности:</b>", label_style))
            table_data = [["Организация", "Должность", "Период работы", "Доп. инфо"]]
            
            for record in work_experience_info["records"]:
                start_date_obj = record.get('date_in')
                end_date_obj = record.get('date_out')

                start_date_str = start_date_obj.strftime("%d.%m.%Y") if start_date_obj else 'N/A'
                end_date_str = end_date_obj.strftime("%d.%m.%Y") if end_date_obj else 'по н.в.'
                
                period_str = f"{start_date_str} - {end_date_str}"
                
                org_details = Paragraph(record.get('organization', 'N/A'), normal_style_left)
                pos_details = Paragraph(record.get('position', 'N/A'), normal_style_left)

                additional_info_str = record.get('raw_text', '')
                additional_info_paragraph = Paragraph(additional_info_str, styles['Normal'])

                table_data.append([
                    org_details,
                    pos_details,
                    Paragraph(period_str, styles['Normal']),
                    additional_info_paragraph
                ])
            
            if len(table_data) > 1:
                work_table = Table(table_data, colWidths=[5*cm, 4*cm, 3.5*cm, 4.5*cm])
                work_table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke),
                    ('ALIGN',(0,0),(-1,-1),'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), bold_font_name),
                    ('FONTNAME', (0,1), (-1,-1), font_name),
                    ('FONTSIZE', (0,0), (-1,-1), 10),
                    ('BOTTOMPADDING', (0,0), (-1,0), 10),
                    ('BACKGROUND', (0,1), (-1,-1), colors.beige),
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ]))
                elements.append(work_table)
            else:
                elements.append(Paragraph("Записи о периодах трудовой деятельности отсутствуют.", normal_style_left))

    benefits = case_details.get("benefits")
    if benefits:
        elements.append(Paragraph("3.3. Заявленные льготы", subsection_title_style))
        for benefit in benefits:
            elements.append(Paragraph(f"• {benefit}", styles['Bullet']))
    
    submitted_documents = case_details.get("submitted_documents")
    if submitted_documents:
        elements.append(Paragraph("3.4. Перечень представленных документов", subsection_title_style))
        for doc_id in submitted_documents:
            doc_name = _get_document_display_name(doc_id, doc_requirements_config, benefit_type_id)
            elements.append(Paragraph(f"• {doc_name} (ID: {doc_id})", styles['Bullet']))

    if case_details.get("has_incorrect_document"):
        elements.append(Paragraph("3.5. Отметка о некорректно оформленных документах", subsection_title_style))
        elements.append(Paragraph("Заявителем отмечено наличие некорректно оформленных документов.", normal_style_left))

    other_documents_data = case_details.get("other_documents_extracted_data")
    if other_documents_data:
        elements.append(Paragraph("3.6. Сведения из дополнительно загруженных документов (по результатам OCR)", subsection_title_style))
        for i, ocr_doc in enumerate(other_documents_data):
            elements.append(Paragraph(f"<b>Документ {i+1}:</b>", label_style))
            doc_type_display = ocr_doc.get("standardized_document_type") or ocr_doc.get("identified_document_type") or "Тип не определен"
            elements.append(Paragraph(f"  <i>Тип документа (определенный системой):</i> {doc_type_display}", normal_style_left))
            extracted_fields = ocr_doc.get("extracted_fields")
            if extracted_fields and isinstance(extracted_fields, dict):
                elements.append(Paragraph("  <i>Ключевые извлеченные поля (обезличенные):</i>", normal_style_left))
                for key, val in extracted_fields.items():
                    val_display = "[СКРЫТО]" if isinstance(val, str) and len(val) > 3 else str(val)
                    elements.append(Paragraph(f"    - {key}: {val_display}", normal_style_left))
            multimodal_assessment = ocr_doc.get("multimodal_assessment")
            if multimodal_assessment:
                 elements.append(Paragraph(f"  <i>Оценка документа системой:</i> {multimodal_assessment}", normal_style_left))
        elements.append(Spacer(1, 0.2*cm))

    elements.append(Spacer(1, 0.5*cm))

    elements.append(Paragraph("4. Результаты автоматизированного анализа и решение", section_title_style))
    final_status = case_details.get("final_status", "Статус не определен")
    status_display = final_status
    if final_status == "СООТВЕТСТВУЕТ":
        status_display = "Право на меры поддержки подтверждено"
    elif final_status == "НЕ СООТВЕТСТВУЕТ":
        status_display = "В праве на меры поддержки отказано (условия не выполнены)"
    elif final_status == "PROCESSING":
        status_display = "Дело находится в обработке"
    elif final_status == "ERROR_PROCESSING":
        status_display = "Ошибка при обработке дела"
        
    elements.append(Paragraph("4.1. Итоговое решение системы", subsection_title_style))
    elements.append(Paragraph(f"<b>{status_display}</b>", normal_style_justify))

    elements.append(Paragraph("4.2. Обоснование решения", subsection_title_style))
    explanation_md = case_details.get("final_explanation", "Обоснование отсутствует.")
    explanation_html_for_pdf = _convert_markdown_to_html_for_reportlab(explanation_md)
    elements.append(Paragraph(explanation_html_for_pdf, normal_style_justify))

    rag_confidence = case_details.get("rag_confidence")
    if rag_confidence is not None:
        elements.append(Paragraph("4.3. Степень уверенности системы в принятом решении", subsection_title_style))
        elements.append(Paragraph(f"{rag_confidence*100:.1f}%", normal_style_justify))
    elements.append(Spacer(1, 0.5*cm))
    
    errors_list = case_details.get("errors", [])
    if errors_list:
        elements.append(Paragraph("5. Выявленные ошибки/несоответствия", section_title_style))
        for err_idx, error_item in enumerate(errors_list):
            elements.append(Paragraph(f"<b>Ошибка {err_idx + 1}:</b>", subsection_title_style))
            elements.append(Paragraph(f"  Код: {error_item.get('code', 'N/A')}", normal_style_left))
            elements.append(Paragraph(f"  Описание: {error_item.get('description', 'N/A')}", normal_style_left))
            if error_item.get('law'):
                elements.append(Paragraph(f"  Основание (закон): {error_item.get('law')}", normal_style_left))
            if error_item.get('recommendation'):
                elements.append(Paragraph(f"  Рекомендация: {error_item.get('recommendation')}", normal_style_left))
        elements.append(Spacer(1, 0.5*cm))

    elements.append(Paragraph("Сформировано автоматизированной системой поддержки участников СВО 'SVO-AI'.", footer_style))
    elements.append(Paragraph("Данное решение носит предварительный характер.", ParagraphStyle(name='FooterDisclaimer', parent=footer_style, fontSize=9)))

    def add_page_numbers(canvas, doc_template):
        canvas.saveState()
        canvas.setFont(font_name, 9)
        page_num_text = f"Страница {doc_template.page}"
        canvas.drawCentredString(A4[0]/2, 1*cm, page_num_text)
        canvas.restoreState()

    doc.build(elements, onFirstPage=add_page_numbers, onLaterPages=add_page_numbers)
    buffer.seek(0)
    return buffer

def _generate_docx_report(
    case_details: Dict[str, Any],
    benefit_types_list_config: List[BenefitTypeInfo],
    doc_requirements_config: Dict[str, BenefitTypeDocuments]
) -> io.BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    case_id = case_details.get("id", "N/A")
    current_date_str = datetime.now().strftime("%d.%m.%Y")
    personal_data = case_details.get("personal_data", {})
    masked_personal_data = mask_personal_data(personal_data if personal_data else {})

    title = doc.add_heading(f"ЗАКЛЮЧЕНИЕ № {case_id}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"по запросу о мерах социальной поддержки", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_p = doc.add_paragraph(f"Дата формирования: {current_date_str}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()

    doc.add_heading("1. Сведения о заявителе (обезличенные)", level=2)
    p = doc.add_paragraph()
    p.add_run("Фамилия, имя, отчество (при наличии): ").bold = True
    p.add_run(masked_personal_data['full_name'])
    p = doc.add_paragraph()
    p.add_run("Дата рождения: ").bold = True
    p.add_run(masked_personal_data['birth_date'])
    p = doc.add_paragraph()
    p.add_run("СНИЛС: ").bold = True
    p.add_run(masked_personal_data['snils'])
    p = doc.add_paragraph()
    p.add_run("Пол: ").bold = True
    p.add_run(masked_personal_data['gender'])
    p = doc.add_paragraph()
    p.add_run("Гражданство: ").bold = True
    p.add_run(masked_personal_data['citizenship'])
    
    if masked_personal_data.get("name_change_info") and masked_personal_data["name_change_info"].get('old_full_name') != "[ПРЕЖНЕЕ ФИО СКРЫТО]":
        p = doc.add_paragraph()
        p.add_run("Сведения о ранее измененном ФИО:").bold = True
        sub_p = doc.add_paragraph(f"  Прежнее ФИО: {masked_personal_data['name_change_info']['old_full_name']}", style='ListBullet')
        sub_p.paragraph_format.left_indent = Cm(0.5)
        sub_p = doc.add_paragraph(f"  Дата изменения: {masked_personal_data['name_change_info']['date_changed']}", style='ListBullet')
        sub_p.paragraph_format.left_indent = Cm(0.5)
        
    p = doc.add_paragraph()
    p.add_run("Количество заявленных иждивенцев: ").bold = True
    p.add_run(masked_personal_data['dependents'])
    doc.add_paragraph()

    doc.add_heading("2. Запрашиваемый вид поддержки", level=2)
    benefit_type_id = case_details.get("pension_type", "Не указан")
    benefit_type_name = _get_benefit_type_display_name(benefit_type_id, benefit_types_list_config)
    doc.add_paragraph(benefit_type_name, style='Normal')
    doc.add_paragraph()

    doc.add_heading("3. Представленные сведения и документы", level=2)
    
    disability_info = case_details.get("disability")
    if disability_info:
        doc.add_heading("3.1. Сведения об инвалидности", level=3)
        dis_date = disability_info.get('date')
        dis_date_str = dis_date.strftime("%d.%m.%Y") if isinstance(dis_date, datetime) or isinstance(dis_date, date) else str(dis_date or 'Не указана')
        p = doc.add_paragraph()
        p.add_run("Дата установления: ").bold = True
        p.add_run(dis_date_str)
        p = doc.add_paragraph()
        p.add_run("Номер справки МСЭ: ").bold = True
        p.add_run(disability_info.get('cert_number', 'Не указан'))

    work_experience_info = case_details.get("work_experience")
    pension_points = case_details.get("pension_points")
    if work_experience_info or pension_points is not None:
        doc.add_heading("3.2. Сведения о трудовом стаже", level=3)
        if work_experience_info:
            total_years = work_experience_info.get('calculated_total_years', 'Не указан')
            p = doc.add_paragraph()
            p.add_run("Общий страховой стаж (лет): ").bold = True
            p.add_run(str(total_years))
        if pension_points is not None:
            p = doc.add_paragraph()
            p.add_run("Индивидуальный пенсионный коэффициент (ИПК): ").bold = True
            p.add_run(str(pension_points))

        if work_experience_info and work_experience_info.get("records"):
            p = doc.add_paragraph()
            p.add_run("Периоды трудовой деятельности:").bold = True
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Организация / Должность'
            hdr_cells[1].text = 'Период работы'
            hdr_cells[2].text = 'Дополнительная информация'

            for record in work_experience_info["records"]:
                row_cells = table.add_row().cells
                
                org_text = record.get('organization', 'N/A')
                pos_text = record.get('position', 'N/A')
                cell1_paragraph = row_cells[0].paragraphs[0]
                cell1_paragraph.add_run(f"{org_text}\n").bold = True
                cell1_paragraph.add_run(pos_text)

                start_date_obj = record.get('date_in')
                end_date_obj = record.get('date_out')
                start_date_str = start_date_obj.strftime("%d.%m.%Y") if start_date_obj else 'N/A'
                end_date_str = end_date_obj.strftime("%d.%m.%Y") if end_date_obj else 'по н.в.'
                row_cells[1].text = f"{start_date_str} - {end_date_str}"
                
                row_cells[2].text = record.get('raw_text', '')

            for column in table.columns:
                column.autofit = True

    benefits = case_details.get("benefits")
    if benefits:
        doc.add_heading("3.3. Заявленные льготы", level=3)
        for benefit in benefits:
            doc.add_paragraph(str(benefit), style='ListBullet')
    
    submitted_documents = case_details.get("submitted_documents")
    if submitted_documents:
        doc.add_heading("3.4. Перечень представленных документов", level=3)
        for doc_id in submitted_documents:
            doc_name = _get_document_display_name(doc_id, doc_requirements_config, benefit_type_id)
            doc.add_paragraph(f"{doc_name} (ID: {doc_id})", style='ListBullet')

    if case_details.get("has_incorrect_document"):
        doc.add_heading("3.5. Отметка о некорректно оформленных документах", level=3)
        doc.add_paragraph("Заявителем отмечено наличие некорректно оформленных документов.", style='Normal')

    other_documents_data = case_details.get("other_documents_extracted_data")
    if other_documents_data:
        doc.add_heading("3.6. Сведения из дополнительно загруженных документов (по результатам OCR)", level=3)
        for i, ocr_doc in enumerate(other_documents_data):
            p = doc.add_paragraph()
            p.add_run(f"Документ {i+1}:").bold = True
            doc_type_display = ocr_doc.get("standardized_document_type") or ocr_doc.get("identified_document_type") or "Тип не определен"
            sub_p = doc.add_paragraph(f"  Тип документа (определенный системой): {doc_type_display}", style='Normal')
            sub_p.paragraph_format.left_indent = Cm(0.5)
            extracted_fields = ocr_doc.get("extracted_fields")
            if extracted_fields and isinstance(extracted_fields, dict):
                sub_p = doc.add_paragraph("  Ключевые извлеченные поля (обезличенные):", style='Normal')
                sub_p.paragraph_format.left_indent = Cm(0.5)
                for key, val in extracted_fields.items():
                    val_display = "[СКРЫТО]" if isinstance(val, str) and len(val) > 3 else str(val)
                    field_p = doc.add_paragraph(f"    - {key}: {val_display}", style='ListBullet')
                    field_p.paragraph_format.left_indent = Cm(1.0)
            multimodal_assessment = ocr_doc.get("multimodal_assessment")
            if multimodal_assessment:
                 sub_p = doc.add_paragraph(f"  Оценка документа системой: {multimodal_assessment}", style='Normal')
                 sub_p.paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph()

    doc.add_heading("4. Результаты автоматизированного анализа и решение", level=2)
    final_status = case_details.get("final_status", "Статус не определен")
    status_display = final_status
    if final_status == "СООТВЕТСТВУЕТ":
        status_display = "Право на меры поддержки подтверждено"
    elif final_status == "НЕ СООТВЕТСТВУЕТ":
        status_display = "В праве на меры поддержки отказано (условия не выполнены)"
    elif final_status == "PROCESSING":
        status_display = "Дело находится в обработке"
    elif final_status == "ERROR_PROCESSING":
        status_display = "Ошибка при обработке дела"
        
    doc.add_heading("4.1. Итоговое решение системы", level=3)
    p = doc.add_paragraph()
    p.add_run(status_display).bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_heading("4.2. Обоснование решения", level=3)
    explanation = case_details.get("final_explanation", "Обоснование отсутствует.")
    doc.add_paragraph(explanation, style='Normal').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rag_confidence = case_details.get("rag_confidence")
    if rag_confidence is not None:
        doc.add_heading("4.3. Степень уверенности системы в принятом решении", level=3)
        doc.add_paragraph(f"{rag_confidence*100:.1f}%", style='Normal').alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph()
    
    errors_list = case_details.get("errors", [])
    if errors_list:
        doc.add_heading("5. Выявленные ошибки/несоответствия", level=2)
        for err_idx, error_item in enumerate(errors_list):
            p_err_title = doc.add_paragraph()
            p_err_title.add_run(f"Ошибка {err_idx + 1}:").bold = True
            
            doc.add_paragraph(f"  Код: {error_item.get('code', 'N/A')}", style='Normal').paragraph_format.left_indent = Cm(0.5)
            doc.add_paragraph(f"  Описание: {error_item.get('description', 'N/A')}", style='Normal').paragraph_format.left_indent = Cm(0.5)
            if error_item.get('law'):
                doc.add_paragraph(f"  Основание (закон): {error_item.get('law')}", style='Normal').paragraph_format.left_indent = Cm(0.5)
            if error_item.get('recommendation'):
                doc.add_paragraph(f"  Рекомендация: {error_item.get('recommendation')}", style='Normal').paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    doc.add_paragraph()
    footer_p1 = doc.add_paragraph("Сформировано автоматизированной системой поддержки участников СВО 'SVO-AI'.")
    footer_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p1.runs: run.font.size = Pt(10)
    
    footer_p2 = doc.add_paragraph("Данное решение носит предварительный характер.")
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p2.runs: run.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

async def generate_document(
    case_details: Dict[str, Any],
    benefit_types_list_config: List[BenefitTypeInfo],
    doc_requirements_config: Dict[str, BenefitTypeDocuments],
    document_format: DocumentFormat
) -> Tuple[io.BytesIO, str, str]:
    doc_format_value = document_format.value

    if doc_format_value == DocumentFormat.pdf.value:
        buffer = await asyncio.to_thread(
            _generate_pdf_report, 
            case_details, benefit_types_list_config, doc_requirements_config
        )
        filename = f"benefit_decision_{case_details.get('id', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.pdf"
        mimetype = "application/pdf"
    elif doc_format_value == DocumentFormat.docx.value:
        buffer = await asyncio.to_thread(
            _generate_docx_report, 
            case_details, benefit_types_list_config, doc_requirements_config
        )
        filename = f"benefit_decision_{case_details.get('id', 'unknown')}_{datetime.now().strftime('%Y%m%d')}.docx"
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        raise ValueError(f"Unsupported document format: {doc_format_value}")

    return buffer, filename, mimetype